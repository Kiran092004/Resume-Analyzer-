import os
import io
import base64
import streamlit as st
from dotenv import load_dotenv
from PIL import Image
import pdf2image
import PyPDF2
import google.generativeai as genai
import mysql.connector
from mysql.connector import Error
from datetime import datetime
import re
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from docx.enum.text import WD_ALIGN_PARAGRAPH


load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

st.set_page_config(page_title="ATS Resume Analyser", layout="wide")


if "show_login" not in st.session_state:
    st.session_state.show_login = False
if "user" not in st.session_state:
    st.session_state.user = None
if "attempts" not in st.session_state:
    st.session_state.attempts = 0  

#  LOGIN BUTTON 
col1, col2 = st.columns([9, 1])
with col2:
    if st.button("🔐 Login"):
        st.session_state.show_login = True


if st.session_state.show_login and not st.session_state.user:
    with st.container():
        st.markdown("""
            <div style="display:flex; justify-content:center; align-items:center; height:60vh;">
                <div style="background:rgba(255,255,255,0.1); padding:2rem; border-radius:12px; width:350px; text-align:center;">
                    <h3 style="color:#00e0ff;"> Login Required</h3>
            """, unsafe_allow_html=True)

        username = st.text_input("Username", key="login_user")
        password = st.text_input("Password", type="password", key="login_pass")
        login_btn = st.button("Login Now")

        if login_btn:
            if username == "admin" and password == "1234":
                st.success(" Login successful!")
                st.session_state.user = username
                st.session_state.show_login = False
            else:
                st.error(" Invalid username or password")

        st.markdown("</div></div>", unsafe_allow_html=True)

#  CSS
st.markdown("""
    <style>
    body { font-family: 'Inter', sans-serif; }
    .main-title { text-align: center; font-size: 2.5rem; font-weight: 700; color: #00e0ff; margin-bottom: 0.5rem; }
    .sub-title { text-align: center; font-size: 1.2rem; color: #cccccc; margin-bottom: 2rem; }
    .glass-card { background: rgba(255, 255, 255, 0.08); border-radius: 16px; padding: 1.5rem; backdrop-filter: blur(10px); border: 1px solid rgba(255, 255, 255, 0.15); margin-bottom: 1rem; }
    .stButton>button { background: linear-gradient(135deg, #00e0ff, #0077ff); color: white; font-weight: 600; border-radius: 12px; padding: 0.6rem 1.2rem; border: none; transition: all 0.3s ease-in-out; }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 4px 15px rgba(0, 224, 255, 0.5); }
    .translucent-output { background: rgba(255, 255, 255, 0.1); padding: 1rem; border-radius: 12px; color: white; font-size: 1rem; line-height: 1.5; }
    </style>
""", unsafe_allow_html=True)


#  DATABASE SETUP 
def get_db_connection():
    try:
        conn = mysql.connector.connect(
            host=os.getenv("MYSQL_HOST"),
            user=os.getenv("MYSQL_USER"),
            password=os.getenv("MYSQL_PASSWORD"),
            database=os.getenv("MYSQL_DATABASE")
        )
        return conn
    except Error as e:
        st.error(f"Database connection error: {e}")
        return None

def init_db():
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS analysis (
                id INT AUTO_INCREMENT PRIMARY KEY,
                timestamp DATETIME,
                resume_name VARCHAR(255),
                job_desc TEXT,
                score INT
            )
        """)
        conn.commit()
        cursor.close()
        conn.close()

def store_analysis(resume_name: str, job_desc: str, score: int):
    conn = get_db_connection()
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO analysis (timestamp, resume_name, job_desc, score)
            VALUES (NOW(), %s, %s, %s)
        """
        cursor.execute(query, (resume_name, job_desc[:200], score))
        conn.commit()
        cursor.close()
        conn.close()


def extract_score(text: str):
    match = re.search(r'(\d{1,3})\s*/\s*100', text)
    if match: return int(match.group(1))
    match = re.search(r'(\d{1,3})\s*%', text)
    if match: return int(match.group(1))
    return None

init_db()

#PDF Process
def input_pdf_setup(uploaded_file):
    pdf_bytes = uploaded_file.read()
    try:
        images = pdf2image.convert_from_bytes(pdf_bytes)
        first_page = images[0]
        img_byte_arr = io.BytesIO()
        first_page.save(img_byte_arr, format="JPEG")
        encoded_image = base64.b64encode(img_byte_arr.getvalue()).decode("utf-8")
        return {"inline_data": {"mime_type": "image/jpeg", "data": encoded_image}}
    except Exception:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        first_page_text = reader.pages[0].extract_text() or ""
        return {"text": first_page_text}


def get_gemini_response(prompt, resume_part=None, job_desc=None):
    model = genai.GenerativeModel('models/gemini-2.5-flash')
    parts = [{"text": prompt}]
    if resume_part:
        if isinstance(resume_part, str):
            parts.append({"text": resume_part})
        else:
            parts.append(resume_part)
    if job_desc:
        parts.append({"text": job_desc})
    response = model.generate_content(parts)
    return response.text

# -------------------- MAIN UI --------------------
st.markdown("<h1 class='main-title'>💼 ATS Resume Analyser</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-title'>Analyse your resume against job descriptions, check grammar, get interview questions & more.</p>", unsafe_allow_html=True)

if not st.session_state.user:
    st.info(f"Free Attempts Left: {max(0, 3 - st.session_state.attempts)}")

with st.container():
    with st.expander(" Upload Resume & Job Description", expanded=True):
        input_text = st.text_area("Paste Job Description", height=150, placeholder="Paste the job description here...")
        uploaded_file = st.file_uploader("Upload Resume (PDF only)", type=["pdf"])
        if uploaded_file:
            st.markdown(f"<p style='color:#00e0ff;'>📄 {uploaded_file.name}</p>", unsafe_allow_html=True)


if st.session_state.attempts < 3 or st.session_state.user:
    st.markdown("### 🔹 Select an Action")
    col1, col2, col3 = st.columns(3)
    with col1: submit_match = st.button(" ATS Match")
    with col2: submit_questions = st.button(" Interview Questions")
    with col3: submit_grammar = st.button(" Grammar Check")
    col4, col5, col6 = st.columns(3)
    with col4: submit_jobs = st.button(" Job Links")
    with col5: submit_missing = st.button(" Missing Skills")
    with col6: submit_new_resume = st.button("Improve Resume")
else:
    st.warning(" You have used 3 free attempts. Please  login to unlock unlimited ATS features.")

# Prompts
prompt_match = "You are a skilled ATS scanner. Provide ATS score out of 100 with analysis."
prompt_questions = "Generate top 5 interview questions for this role."
prompt_grammar = "Check resume grammar issues, suggest corrections."
prompt_links = "Suggest top job portal links for this job description."
prompt_missing = "List missing keywords/skills from resume compared to job description."


if 'submit_match' in locals() and submit_match:
    st.session_state.attempts += 1
    if uploaded_file and input_text:
        with st.spinner("Calculating ATS match..."):
            pdf_content = input_pdf_setup(uploaded_file)
            result = get_gemini_response(prompt_match, pdf_content, input_text)
            st.markdown(f"<div class='translucent-output'>{result}</div>", unsafe_allow_html=True)
            score_val = extract_score(result)
            if score_val is not None:
                store_analysis(uploaded_file.name, input_text, score_val)
                st.success(f" ATS Score: {score_val}%")
                fig = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=score_val,
                    gauge={'axis': {'range': [0, 100]},
                           'bar': {'color': "#00e0ff"},
                           'steps': [{'range': [0,50],'color':'red'},
                                     {'range':[50,75],'color':'yellow'},
                                     {'range':[75,100],'color':'green'}]},
                    title={'text': "ATS Match Score"}
                ))
                st.plotly_chart(fig)
    else:
        st.warning(" Please upload a resume and paste job description.")


if 'submit_questions' in locals() and submit_questions and input_text:
    st.session_state.attempts += 1
    with st.spinner("Fetching interview questions..."):
        result = get_gemini_response(prompt_questions, {}, input_text)
        st.markdown(f"<div class='translucent-output'>{result}</div>", unsafe_allow_html=True)

if 'submit_grammar' in locals() and submit_grammar and uploaded_file:
    st.session_state.attempts += 1
    with st.spinner("Checking grammar..."):
        pdf_content = input_pdf_setup(uploaded_file)
        result = get_gemini_response(prompt_grammar, pdf_content, "")
        st.markdown(f"<div class='translucent-output'>{result}</div>", unsafe_allow_html=True)

if 'submit_jobs' in locals() and submit_jobs and input_text:
    st.session_state.attempts += 1
    with st.spinner("Fetching job links..."):
        result = get_gemini_response(prompt_links, {}, input_text)
        st.markdown(f"<div class='translucent-output'>{result}</div>", unsafe_allow_html=True)

if 'submit_missing' in locals() and submit_missing and uploaded_file and input_text:
    st.session_state.attempts += 1
    with st.spinner("Finding missing skills..."):
        pdf_content = input_pdf_setup(uploaded_file)
        result = get_gemini_response(prompt_missing, pdf_content, input_text)
        st.markdown(f"<div class='translucent-output'>{result}</div>", unsafe_allow_html=True)
        missing_skills = len([line for line in result.split("\n") if line.strip()])
        total_skills = missing_skills + 10
        matched_skills = total_skills - missing_skills
        fig = px.pie(values=[matched_skills, missing_skills],
                     names=["Matched Skills", "Missing Skills"],
                     color_discrete_map={"Matched Skills": "green", "Missing Skills": "red"})
        st.plotly_chart(fig)


if 'submit_new_resume' in locals() and submit_new_resume and uploaded_file and input_text:
    st.session_state.attempts += 1
    with st.spinner(" Crafting a professional, ATS-optimized resume..."):
        pdf_content = input_pdf_setup(uploaded_file)
        resume_prompt = """Rewrite this resume in a modern, ATS-optimized format.
        It should be have good Template.
        Include clear sections: Name, Title, Summary, Experience, Education, Skills, and Certifications (if any).
        Focus on readability, concise phrasing, and keyword density relevant to the job description."""
        
        improved_resume_text = get_gemini_response(resume_prompt, pdf_content, input_text)
        missing_result = get_gemini_response(prompt_missing, pdf_content, input_text)
        missing_skills = [s.strip("•- ") for s in missing_result.split("\n") if s.strip()]

        
        doc = Document()
        header = doc.add_paragraph()
        header_run = header.add_run("IMPROVED RESUME")
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run.bold = True
        header_run.font.size = Pt(20)
        header_run.font.color.rgb = RGBColor(0, 119, 255)

        sub = doc.add_paragraph("Optimized and reformatted for better ATS compatibility and recruiter readability.")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True
        sub.runs[0].font.size = Pt(11)
        doc.add_paragraph("")

        p = doc.add_paragraph()
        run = p.add_run("_" * 100)
        run.font.color.rgb = RGBColor(0, 224, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        for line in improved_resume_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if re.match(r"^[A-Z ]+:$", line.strip(), re.I):
                p = doc.add_paragraph()
                run = p.add_run(line.strip(":").title())
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 184, 255)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
            elif line.startswith(("-", "•")):
                p = doc.add_paragraph(line.strip("-• "), style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.25

        doc.add_page_break()
        ms_title = doc.add_paragraph("⚡ Suggested Missing Skills")
        ms_title.runs[0].bold = True
        ms_title.runs[0].font.size = Pt(14)
        ms_title.runs[0].font.color.rgb = RGBColor(255, 180, 0)

        for skill in missing_skills:
            p = doc.add_paragraph(f"• {skill}", style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

        buffer_docx = BytesIO()
        doc.save(buffer_docx)
        buffer_docx.seek(0)

      
        buffer_pdf = BytesIO()
        pdf = SimpleDocTemplate(buffer_pdf, pagesize=A4, leftMargin=1*inch, rightMargin=1*inch, topMargin=0.8*inch, bottomMargin=0.8*inch)
        styles = getSampleStyleSheet()
        style_normal = styles["Normal"]
        style_heading = styles["Heading2"]
        style_heading.textColor = "#0077ff"
        style_normal.fontSize = 11
        style_normal.leading = 14

        story = []
        story.append(Paragraph("<b><font size=16 color='#0077ff'>IMPROVED RESUME</font></b>", style_heading))
        story.append(Spacer(1, 6))
        story.append(Paragraph("Optimized for ATS and recruiter readability.", style_normal))
        story.append(Spacer(1, 12))

        for line in improved_resume_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if re.match(r"^[A-Z ]+:$", line.strip(), re.I):
                story.append(Spacer(1, 6))
                story.append(Paragraph(f"<b><font color='#00b8ff'>{line.strip(':').title()}</font></b>", style_heading))
                story.append(Spacer(1, 4))
            elif line.startswith(("-", "•")):
                story.append(Paragraph(f"• {line.strip('-• ')}", style_normal))
            else:
                story.append(Paragraph(line, style_normal))

        story.append(Spacer(1, 16))
        story.append(Paragraph("<b><font color='#ffb400'>⚡ Suggested Missing Skills</font></b>", style_heading))
        story.append(Spacer(1, 8))
        if missing_skills:
            skills_data = [[f"• {s}"] for s in missing_skills]
            table = Table(skills_data, colWidths=[5.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.whitesmoke),
                ('TEXTCOLOR', (0,0), (-1,-1), colors.black),
                ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
                ('FONTSIZE', (0,0), (-1,-1), 10),
                ('LEFTPADDING', (0,0), (-1,-1), 6),
                ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ]))
            story.append(table)
        pdf.build(story)
        buffer_pdf.seek(0)

        st.success("Improved Resume Ready with Missing Skills")
        colA, colB = st.columns(2)
        with colA:
            st.download_button(" Download Word", data=buffer_docx,
                               file_name="Improved_Resume.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with colB:
            st.download_button(" Download PDF", data=buffer_pdf,
                               file_name="Improved_Resume.pdf",
                               mime="application/pdf")
